import os
import time
import streamlit as st
import json
from docx import Document
from typing import Dict
from datetime import datetime

# Import improved extraction module
from utils.extraction import extract_field_information, clear_old_sessions, map_extraction_to_database

# Import gestion utilities for database operations
from utils.gestion import save_to_database

# Simple database operations
class SimpleDBManager:
    def save_extraction_to_db(self, results, run_id):
        """Simple database storage simulation"""
        os.makedirs("db", exist_ok=True)
        with open(f"db/extraction_{run_id}.json", "w", encoding="utf-8") as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        return run_id

def create_output_document(results: Dict[str, str], output_path: str) -> None:
    """
    Generate a Word document with extracted information.
    """
    doc = Document()
    doc.add_heading("Fiche d'Appel d'Offres Marocain", level=1)
    
    for field, ans in results.items():
        if field not in ["Error", "Status"]:  # Skip error/status fields
            doc.add_heading(field, level=2)
            doc.add_paragraph(ans)
    
    doc.save(output_path)

def main():
    st.title("Génération de la fiche de dépouillement")
    st.markdown("""
    Téléversez les documents d'appel d'offres pour extraire automatiquement 
    les informations importantes et générer une fiche de dépouillement.
    """)
    
    # Clean up old sessions
    clear_old_sessions()
    
    # Check if we already have processed documents
    if st.session_state.get('document_processed', False):
        st.success("Documents déjà traités et prêts.")
        
        # Display previous extraction results
        if 'document_data' in st.session_state and st.session_state.document_data:
            with st.expander("Résultats d'extraction"):
                for field, value in st.session_state.document_data.items():
                    if field not in ["Error", "Status"]:
                        st.markdown(f"**{field}**: {value}")
            
            # Enhanced action buttons
            col1, col2, col3 = st.columns(3)
            
            with col1:
                # Create and provide Word document download
                if st.button("📄 Télécharger la fiche Word"):
                    try:
                        os.makedirs("output", exist_ok=True)
                        outp = "output/Fiche_Appel_Offres.docx"
                        create_output_document(st.session_state.document_data, outp)
                        
                        with open(outp, "rb") as f:
                            st.download_button(
                                "📥 Télécharger le fichier Word",
                                data=f,
                                file_name="Fiche_Appel_Offres.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    except Exception as e:
                        st.error(f"Erreur lors de la création du fichier Word: {e}")
            
            with col2:
                # NEW: Save to database button
                if st.button("💾 Sauvegarder vers base de données"):
                    try:
                        # Map extraction results to database format
                        db_record = map_extraction_to_database(st.session_state.document_data)
                        
                        # Convert database format to form format expected by save_to_database()
                        form_data = {
                            "reference_ao": db_record.get("Référence AO", ""),
                            "objet": db_record.get("Objet de l'appel d'offre", ""),
                            "organisme_emetteur": db_record.get("Organisme émetteur", ""),
                            "region": db_record.get("Région / Ville"),
                            "secteur": db_record.get("Secteur"),
                            "montant_estime": db_record.get("Montant estimé (MAD)"),
                            "caution": db_record.get("Caution demandée (MAD)"),
                            "date_publication": db_record.get("Date de publication"),
                            "go_no_go": db_record.get("GO / NO GO"),
                            "statut": db_record.get("Statut"),
                            "motif_rejet": db_record.get("Motif de rejet"),
                            "complexite": db_record.get("Complexité perçue (1-5)"),
                            "type_mission": db_record.get("Type de mission"),
                            "responsable": db_record.get("Responsable"),
                            "montant_offert": db_record.get("Montant offert (MAD)"),
                            "duree_marche": db_record.get("Durée du marché (mois)"),
                            "nb_concurrents": db_record.get("Nombre de concurrents (si dispo)"),
                            "date_soumission": db_record.get("Date de soumission"),
                            "date_decision": db_record.get("Date de décision"),
                            "score_technique": db_record.get("Score technique (si dispo)"),
                            "lien_dossier": db_record.get("Lien vers dossier"),
                            "temps_traitement": db_record.get("Temps de traitement (jours)"),
                            "ecart_montant": db_record.get("Écart montant (%)"),
                            "score_strategique": None  # Will be calculated
                        }
                        
                        # Validate required fields
                        if not form_data.get("reference_ao"):
                            st.error("❌ Référence AO manquante - impossible de sauvegarder")
                            return
                        
                        if not form_data.get("organisme_emetteur"):
                            st.error("❌ Organisme émetteur manquant - impossible de sauvegarder")
                            return
                        
                        # Save to database
                        success, message = save_to_database(form_data)
                        
                        if success:
                            st.success(f"✅ {message}")
                            st.info("💡 Vous pouvez maintenant compléter les informations de gestion dans la page 'Gestion des AO'")
                            st.balloons()
                        else:
                            st.error(f"❌ {message}")
                            
                    except Exception as e:
                        st.error(f"❌ Erreur lors de la sauvegarde: {str(e)}")
                        st.exception(e)  # Show full error for debugging
            
            with col3:
                # Option to restart
                if st.button("🔄 Traiter de nouveaux documents"):
                    st.session_state.document_processed = False
                    st.session_state.document_data = {}
                    st.rerun()
            
            # Quick navigation
            st.markdown("---")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("📊 Aller au Dashboard"):
                    st.switch_page("pages/04_dashboard.py")
            with col2:
                if st.button("📝 Aller à la Gestion des AO"):
                    st.switch_page("pages/03_gestion.py")
        
        return
    
    # File uploaders
    col1, col2, col3 = st.columns(3)
    
    # Create dictionary of file uploaders - using simple, clear keys
    upload = {
        "rc": col1.file_uploader("Téléversez le Règlement de Consultation (RC)", type="pdf"),
        "cps": col2.file_uploader("Téléversez le Cahier des Prescriptions Spéciales (CPS)", type="pdf"),
        "avis": col3.file_uploader("Téléversez l'Avis d'Appel d'Offres", type="pdf"),
    }
    
    # Process button
    if st.button("🚀 Lancer l'extraction et Génération", type="primary"):
        # Check if at least one file is uploaded
        if not any(upload.values()):
            st.error("Veuillez téléverser au moins un document.")
            return

        try:
            # Generate a unique run ID
            run_id = f"extraction_{int(time.time())}"
            st.session_state.run_id = run_id
            
            with st.spinner("Traitement des documents en cours..."):
                # Extract information using session-isolated approach
                results = extract_field_information(upload)
                
                # Validate results
                if not results:
                    st.error("L'extraction n'a pas retourné de résultats.")
                    return
                
                if "Error" in results:
                    st.error(f"Erreur d'extraction: {results['Error']}")
                    return
                
                # Store in session state
                st.session_state.document_data = results
                st.session_state.document_processed = True
                
                # Display results
                st.subheader("📋 Informations Extraites")
                
                # Organize results in a nice format
                col1, col2 = st.columns(2)
                
                fields_left = ["Référence", "Objet", "Maître d'Ouvrage", "Date", "Estimation des coûts", "Montant de la caution"]
                fields_right = ["Contact", "Contenu Dossier", "Modalités de retrait", "Offre Financière", "Offre Technique"]
                
                with col1:
                    st.markdown("### 📊 Informations Principales")
                    for field in fields_left:
                        if field in results:
                            st.markdown(f"**{field}**: {results[field]}")
                
                with col2:
                    st.markdown("### 📋 Informations Complémentaires")
                    for field in fields_right:
                        if field in results:
                            st.markdown(f"**{field}**: {results[field]}")

                # Show success message
                st.success("✅ Extraction terminée avec succès!")
                st.info("💡 Utilisez les boutons ci-dessus pour télécharger la fiche ou sauvegarder vers la base de données")
                
                # Auto-scroll to top to show the action buttons
                st.rerun()

        except Exception as e:
            st.error(f"Une erreur s'est produite: {e}")
            st.exception(e)  # Show full traceback for debugging

if __name__ == "__main__":
    main()